from base64 import b64decode
import copy
from io import BytesIO
import re
from typing import List, Tuple

from arches_templating.template_engine.template_engine_factory import TemplateEngineFactory
from arches_templating.template_engine.template_tag_type import TemplateTagType
from arches_templating.template_engine.template_engine import TemplateEngine
from arches_templating.template_engine.template_tag import TemplateTag
import docx
from docx.oxml.ns import qn
from docx.document import Document as _Document
from docx.document import _Body
from docx.oxml import OxmlElement
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.section import _Header

@TemplateEngineFactory.register('docx')
class DocxTemplateEngine(TemplateEngine):
    def extract_regex_matches(self, template) -> List[Tuple]:
        self.doc = docx.Document(template)
        tags = self.iterate_over_container(self.doc)

        return tags

    def iterate_over_container(self, container, parent=None):
        parsed_tags: List[Tuple] = []
        try:
            for section in container.sections:
                parsed_tags += self.iterate_over_container(section.header)
        except AttributeError:
            pass  # this is ok, there are lots of types that do not have a "sections" attribute - skip them and continue

        for block in self.iter_block_items(container):
            if isinstance(block, Paragraph):
                for match in re.findall(self.regex, block.text):
                    parsed_tags.append((match, {"docxBlock": block, "parent": parent}))

            elif isinstance(block, Table):
                row_length = len(block.rows)
                column_length = len(block.columns)
                current_row = 0
                while current_row < row_length:
                    current_column = 0
                    while current_column < column_length:
                        current_cell = block.cell(current_row, current_column)
                        parsed_tags += self.iterate_over_container(current_cell, block)
                        current_column += 1
                    current_row += 1
                pass
        return parsed_tags

    def iter_block_items(self, parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
            element_parent = parent._body
        elif isinstance(parent, _Cell):
            element_parent = parent_elm = parent._tc
        elif isinstance(parent, _Header):
            element_parent = parent_elm = parent._element
        elif isinstance(parent, _Body):
            parent_elm = parent._body
            element_parent = parent
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, element_parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, element_parent)

    def delete_paragraph(paragraph):
        p = paragraph._element
        paragraph_parent = p.getparent()
        if paragraph_parent is not None: 
            paragraph_parent.remove(p)
            p._p = p._element = None
    
    def delete_table(table):
        table._element.getparent().remove(table._element)

    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def replace_tags(self, tags: List[TemplateTag]):
        incomplete = False
        for tag in tags:
            block = tag.optional_keys["docxBlock"]
            if tag.type == TemplateTagType.CONTEXT:
                if tag.has_rows and tag.context_children_template:
                    # render a table
                    parent = tag.context_children_template[-1].optional_keys["parent"]
                    if isinstance(parent, Table):
                    
                        column = 0
                        # this is ugly, but way more efficient than the alternative
                        
                        current_row = parent.add_row()

                        for child in tag.children:
                            if child.type == TemplateTagType.ROWEND:
                                column = -1
                                current_row = parent.add_row()
                            elif child.type == TemplateTagType.VALUE:
                                # grab any borders from the original cell copy them to the new cell.
                                template_block = tag.context_children_template[column].optional_keys["docxBlock"]
                                borders = template_block._parent.get_or_add_tcPr().first_child_found_in("w:tcBorders")

                                for edge in ("start", "top", "end", "bottom", "left", "right", "insideH", "insideV"):
                                    raw_border_tag = "w:{}".format(edge)

                                    # check for tag existnace, if none found, then create one
                                    element = borders.find(qn(raw_border_tag))
                                    cell_borders = current_row.cells[column]._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
                                    if cell_borders is None:
                                        cell_borders = OxmlElement("w:tcBorders")
                                        current_row.cells[column]._tc.get_or_add_tcPr().append(cell_borders)
                                    if element is not None:
                                        cell_borders.append(copy.deepcopy(element))
                                # every cell gets created with (bad) default styling.
                                DocxTemplateEngine.delete_paragraph(current_row.cells[column].paragraphs[0])
                                # copies paragraph styling from the original template cells
                                current_row.cells[column].add_paragraph(
                                    "" if child.value == None else child.value,
                                    copy.deepcopy(_Cell(template_block._parent, parent).paragraphs[0].style),
                                )
                            column += 1

                        if tag.attributes["has_header"] == "true":
                            DocxTemplateEngine.remove_row(parent, parent.rows[1])
                        else:
                            DocxTemplateEngine.remove_row(parent, parent.rows[0])
                    else:
                        incomplete = True
                        all_blocks_in_context = []
                        # get all blocks between context start and end
                        for item in self.iterate_inner_block(block, tag):
                            all_blocks_in_context.append(item)
                        
                        # for each child (row) copy the context section
                        context_index = 0
                        while context_index < tag.context_length:
                            block_index = 0
                            while block_index <= len(all_blocks_in_context) - 1:
                                if block_index == 0:
                                    new_index_attribute = "index=\"{}\" ".format(context_index)
                                    match = re.findall(self.regex, all_blocks_in_context[block_index].text)
                                    try:
                                        tag_length = len(all_blocks_in_context[block_index].text)
                                        attribute_length = len(match[0][2])
    
                                        insert_index = tag_length - attribute_length
                                        new_tag_text = all_blocks_in_context[block_index].text[:insert_index] + new_index_attribute + all_blocks_in_context[block_index].text[insert_index:]

                                    except IndexError:
                                        # OK to fail, means there was an issue getting a tag match.  Bail/abort rather than scream.
                                        
                                        pass

                                    tag.end_tag.optional_keys["docxBlock"].insert_paragraph_before(new_tag_text, all_blocks_in_context[block_index].style)
                                else:
                                    tag.end_tag.optional_keys["docxBlock"]._element.addprevious(copy.deepcopy(all_blocks_in_context[block_index]._element))
                                    
                                block_index += 1
                            context_index += 1
                        for original_block in all_blocks_in_context:
                            DocxTemplateEngine.delete_paragraph(original_block)
                else:
                    replace_tags_result = self.replace_tags(tag.children)
                    incomplete = incomplete or replace_tags_result
                    
                DocxTemplateEngine.delete_paragraph(block)
                DocxTemplateEngine.delete_paragraph(tag.end_tag.optional_keys["docxBlock"])
            elif tag.type == TemplateTagType.VALUE:
                block.text = tag.value
            elif tag.type == TemplateTagType.IMAGE:
                block.text = ""
                run = block.add_run()
                if tag.value:
                    run.add_picture(BytesIO(b64decode(re.sub("data:image/jpeg;base64,", "", tag.value))))
            # do not use if blocks within tables (note as of 6/6/2023)
            elif tag.type == TemplateTagType.IF:
                if tag.render:
                    DocxTemplateEngine.delete_paragraph(block)
                    DocxTemplateEngine.delete_paragraph(tag.end_tag.optional_keys["docxBlock"])
                    self.replace_tags(tag.children)
                else:
                    for item in self.iterate_inner_block(block, tag):
                        DocxTemplateEngine.delete_paragraph(item)
        return incomplete

                
    def iterate_inner_block(self, block, tag):
        found_if_start = False
        found_if_end = False
        for item in self.iter_block_items(block._parent):
            if item._element == block._element:
                found_if_start = True
            if item._element == tag.end_tag.optional_keys["docxBlock"]._element:
                found_if_end = True
                yield item 
            if found_if_start and not found_if_end:
                yield item


    def create_file(self, tags: List[TemplateTag], template):
        bytestream = BytesIO()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        incomplete = self.replace_tags(tags)
        self.doc.save(bytestream)
        return (bytestream, mime, incomplete)
